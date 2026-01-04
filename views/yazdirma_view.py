"""
Kelebek Sınav Sistemi - Yazdırma Merkezi
Salon bazlı dosya üretimi ve öğrenci listesi
"""

import tkinter as tk
from tkinter import ttk, filedialog
import os
import shutil
from datetime import datetime
from typing import Dict
from io import BytesIO

import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from assets.styles import (KelebekTheme, create_card_frame, configure_standard_button,
                           show_message, ScrollableFrame)
from assets.layout import setup_responsive_window
from controllers.database_manager import get_db
from utils import format_sira_label
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from docx import Document


class YazdirmaView:
    """Yazdırma ekranı"""
    
    def __init__(self, window, parent):
        self.window = window
        self.parent = parent
        self.db = get_db()
        
        self.salon_map = {}
        self.salon_data = {}
        self.selected_salon = None
        self.salon_var = tk.StringVar(value="Seçilmedi")
        self.exam_lookup = {}
        
        self.setup_ui()
        self.load_exam_salons()
    
    def setup_ui(self):
        self.window.title(f"{KelebekTheme.ICON_PRINT} O DERS SAATİNDE YAPILACAK TÜM SINAV KAĞITLARINI OTURMA DÜZENİNE GÖRE YAZDIR")
        self.window.geometry("1100x700")
        self.window.config(bg=KelebekTheme.BG_LIGHT)
        setup_responsive_window(self.window)
        
        header = tk.Frame(self.window, bg=KelebekTheme.PRIMARY, height=60)
        header.pack(fill="x")
        header.pack_propagate(False)
        
        tk.Label(
            header,
            text=f"{KelebekTheme.ICON_PRINT} SINAV YAZDIRMA MERKEZİ",
            font=(KelebekTheme.FONT_FAMILY, 20, "bold"),
            fg=KelebekTheme.TEXT_WHITE,
            bg=KelebekTheme.PRIMARY
        ).pack(pady=15)
        
        # Ana konteyner (Split V.)
        main_frame = tk.Frame(self.window, bg=KelebekTheme.BG_LIGHT)
        main_frame.pack(fill="both", expand=True, padx=15, pady=15)
        
        # Sol Panel: Kontroller (width=350)
        left_panel = tk.Frame(main_frame, bg=KelebekTheme.BG_LIGHT, width=350)
        left_panel.pack(side="left", fill="y", padx=(0, 15))
        left_panel.pack_propagate(False)
        
        control_card = create_card_frame(left_panel, "Yazdırma İşlemleri", KelebekTheme.ICON_PRINT)
        control_card.pack(fill="both", expand=True)
        self.create_control_panel(control_card.content)
        
        # Sağ Panel: Liste
        right_panel = tk.Frame(main_frame, bg=KelebekTheme.BG_LIGHT)
        right_panel.pack(side="right", fill="both", expand=True)
        
        list_card = create_card_frame(right_panel, "Sınav Kağıdı Önizleme & Liste", KelebekTheme.ICON_SEARCH)
        list_card.pack(fill="both", expand=True)
        self.create_list_panel(list_card.content)
    
    def create_control_panel(self, parent):
        """Sol panel - Kontroller"""
        # Salon Seçimi
        tk.Label(
            parent,
            text="Salon Seçimi:",
            font=(KelebekTheme.FONT_FAMILY, 10, "bold"),
            bg=KelebekTheme.BG_WHITE
        ).pack(anchor="w", padx=10, pady=(10, 5))
        
        self.salon_combo = ttk.Combobox(
            parent,
            textvariable=self.salon_var,
            state="readonly"
        )
        self.salon_combo.pack(fill="x", padx=10, pady=5)
        self.salon_combo.bind("<<ComboboxSelected>>", self.on_salon_selected)
        
        btn_reload = tk.Button(parent, command=self.load_exam_salons, text="↻ Listeyi Yenile")
        configure_standard_button(btn_reload, "secondary", "↻ Listeyi Yenile")
        btn_reload.pack(fill="x", padx=10, pady=5)
        
        tk.Label(parent, text="-"*30, bg=KelebekTheme.BG_WHITE, fg="#eee").pack(pady=10)
        
        # Bilgi
        tk.Label(
            parent,
            text="Sınav Bilgisi:",
            font=(KelebekTheme.FONT_FAMILY, 10, "bold"),
            bg=KelebekTheme.BG_WHITE
        ).pack(anchor="w", padx=10, pady=(10, 5))
        
        self.exam_info_label = tk.Label(
            parent,
            text="Seçim bekleniyor...",
            font=(KelebekTheme.FONT_FAMILY, 9),
            fg=KelebekTheme.TEXT_DARK,
            bg=KelebekTheme.BG_WHITE,
            justify="left",
            wraplength=300,
            anchor="w"
        )
        self.exam_info_label.pack(fill="x", padx=10, pady=5)
        
        tk.Frame(parent, bg=KelebekTheme.BG_WHITE, height=20).pack() # Spacer

        # Butonlar (Dikey yerleşim)
        btn_salon = tk.Button(parent, command=self.print_selected_salon, wraplength=300)
        configure_standard_button(btn_salon, "primary", "🖨️ SEÇİLİ SALONU YAZDIR\n(Word + Excel Olarak Kaydet)")
        btn_salon.pack(fill="x", padx=10, pady=10)
        
        btn_all = tk.Button(parent, command=self.print_all_available_salons, wraplength=300)
        configure_standard_button(btn_all, "success", "🖨️ TÜM OKULU YAZDIR\n(Tüm Salonlar İçin Klasör Üret)")
        btn_all.pack(fill="x", padx=10, pady=10)

    def create_list_panel(self, parent):
        """Sağ panel - Tablo"""
        tree_frame = tk.Frame(parent, bg=KelebekTheme.BG_WHITE)
        tree_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        columns = ("Sıra", "Ad", "Soyad", "Sınıf/Şube", "Salon", "Gözetmen", "Sınav Bilgisi")
        self.student_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=20)
        
        # Kolon genişlikleri optimize edildi
        widths = [50, 120, 120, 80, 100, 150, 200]
        for col, width in zip(columns, widths):
            self.student_tree.heading(col, text=col)
            self.student_tree.column(col, width=width, anchor="center")
        
        self.student_tree.pack(side="left", fill="both", expand=True)
        
        scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.student_tree.yview)
        self.student_tree.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")
    
    def load_exam_salons(self):
        try:
            exams = self.db.harmanlanmis_sinavlar()
        except Exception as e:
            show_message(self.window, f"Sınavlar yüklenemedi: {e}", "error")
            return
        
        self.salon_map.clear()
        self.salon_data = {}
        self.exam_lookup = {}
        options: list[str] = []
        
        for exam in exams:
            try:
                yerlesim = self.db.yerlesim_getir(exam['id'])
            except Exception:
                continue
            if not yerlesim:
                continue
            gozetmen_map = self._fetch_gozetmen_map(exam['id'])
            annotated = [
                {
                    **yer,
                    'sinav_id': exam['id'],
                    'sinav_adi': exam['sinav_adi'],
                    'sinav_tarihi': exam.get('sinav_tarihi') or '-',
                    'sinav_saati': exam.get('sinav_saati') or '-',
                    'ders_adi': exam['ders_adi']
                }
                for yer in yerlesim
            ]
            self.exam_lookup[exam['id']] = exam
            
            salon_groups = {}
            for yer in annotated:
                salon_groups.setdefault(yer['salon_adi'], []).append(yer)
            
            for salon_name, placement in salon_groups.items():
                salon_id = placement[0]['salon_id']
                entry = {
                    'exam': exam,
                    'exam_id': exam['id'],
                    'placement': placement,
                    'salon': salon_name,
                    'gozetmen': gozetmen_map.get(salon_id, "")
                }
                self.salon_map.setdefault(salon_name, []).append(entry)
        
        options = sorted(self.salon_map.keys())
        self._rebuild_salon_data()
        self.salon_combo['values'] = options
        if options:
            self.salon_combo.current(0)
            self.on_salon_selected()
        else:
            self.salon_combo.set("Seçilmedi")
            self.selected_salon = None
            self.exam_info_label.config(text="Harmanlanmış salon bulunamadı.")
            self.populate_students()

    def _rebuild_salon_data(self):
        self.salon_data = {}
        for salon_name, entries in self.salon_map.items():
            students = []
            gozetmen = ""
            for entry in entries:
                students.extend(entry['placement'])
                if not gozetmen and entry.get('gozetmen'):
                    gozetmen = entry['gozetmen']
            students.sort(key=lambda x: (x['sira_no'], x.get('sinav_adi', '')))
            salon_id = students[0]['salon_id'] if students else None
            self.salon_data[salon_name] = {
                'salon_id': salon_id,
                'students': students,
                'gozetmen': gozetmen
            }
    
    def on_salon_selected(self, event=None):
        salon_name = self.salon_var.get()
        if salon_name not in self.salon_data:
            self.reset_selection_state("Salon seçilmedi.")
            return
        self.selected_salon = salon_name
        self.update_exam_info_label()
        self.populate_students(salon_name)

    def reset_selection_state(self, message=""):
        self.selected_salon = None
        self.exam_info_label.config(text=message or "Salon seçilmedi.")
        self.populate_students()

    def get_exam_info_text(self):
        salon_name = self.selected_salon
        salon = self.salon_data.get(salon_name)
        if not salon or not salon_name:
            return "Salon seçilmedi."
        toplam = len(salon['students'])
        exam_counts = {}
        for yer in salon['students']:
            exam_counts[yer.get('sinav_adi', "Bilinmeyen")] = exam_counts.get(yer.get('sinav_adi', "Bilinmeyen"), 0) + 1
        info_lines = [
            f"Salon: {salon_name}",
            f"Sınav sayısı: {len(exam_counts)} | Öğrenci: {toplam}"
        ]
        preview_items = list(exam_counts.items())[:3]
        if preview_items:
            preview = ", ".join(f"{adi} ({sayi})" for adi, sayi in preview_items)
            if len(exam_counts) > 3:
                preview += ", ..."
            info_lines.append(preview)
        return "\n".join(info_lines)

    def update_exam_info_label(self):
        self.exam_info_label.config(text=self.get_exam_info_text())
    
    def _fetch_gozetmen_map(self, exam_id):
        try:
            atamalar = self.db.gozetmen_atamalari_listele(exam_id)
        except AttributeError:
            return {}
        temp = {}
        for atama in atamalar:
            gorev = "Asıl" if atama['gorev_turu'] == 'asil' else "Yedek"
            temp.setdefault(atama['salon_id'], []).append(f"{atama['ad']} {atama['soyad']} ({gorev})")
        return {sid: ", ".join(names) for sid, names in temp.items()}
    
    def populate_students(self, salon_name=None):
        for item in self.student_tree.get_children():
            self.student_tree.delete(item)
        
        if salon_name:
            self.selected_salon = salon_name
        salon = self.salon_data.get(self.selected_salon)
        if not salon:
            return
        
        gozetmen_text = salon.get('gozetmen', "")
        
        for yer in sorted(salon['students'], key=lambda x: (x.get('sinav_adi', ''), x['sira_no'])):
            seat_label = self._seat_display(yer)
            exam_info = (f"{yer.get('sinav_adi', '-')}"
                         f" ({yer.get('sinav_tarihi', '-')}"
                         f" {yer.get('sinav_saati', '')})")
            self.student_tree.insert("", "end", values=(
                seat_label,
                yer['ad'],
                yer['soyad'],
                f"{yer['sinif']}/{yer['sube']}",
                self.selected_salon,
                gozetmen_text,
                exam_info
            ))
    
    def _refresh_exam(self, exam_id):
        refreshed = self.db.sinav_getir(exam_id)
        if refreshed:
            self.exam_lookup[exam_id] = refreshed
            return refreshed
        return self.exam_lookup.get(exam_id)

    def _ensure_exam_file(self, exam):
        if not exam:
            return None
        if exam.get('id'):
            exam = self._refresh_exam(exam['id']) or exam
        soru = exam.get('soru_dosyasi')
        source_path = None
        if soru:
            source_path = soru.get('dosya_yolu')
            if not source_path or not os.path.isfile(source_path):
                source_path = None
        if not source_path:
            source_path = self._auto_attach_question_file(exam)
        if not source_path:
            show_message(
                self.window,
                f"{exam.get('sinav_adi', 'Sınav')} için soru dosyası bulunamadı!",
                "warning"
            )
            return None
        return source_path

    def _auto_attach_question_file(self, exam):
        ders_id = exam.get('ders_id')
        if not ders_id:
            return None
        try:
            dosyalar = self.db.soru_bankasi_listele(ders_id=ders_id)
        except Exception:
            return None
        for dosya in dosyalar:
            path = dosya.get('dosya_yolu')
            if not path or not os.path.isfile(path):
                continue
            soru_info = {
                'id': dosya['id'],
                'dosya_adi': dosya['dosya_adi'],
                'dosya_yolu': dosya['dosya_yolu'],
                'orjinal_dosya': dosya.get('orjinal_dosya')
            }
            exam['soru_dosyasi'] = soru_info
            if exam.get('id'):
                try:
                    self.db.sinav_soru_dosyasi_guncelle(exam['id'], dosya['id'])
                except Exception:
                    pass
            show_message(
                self.window,
                f"Bu sınav için kayıtlı soru dosyası bulunamadı. Dersin en güncel dosyası "
                f"otomatik atandı: {dosya['dosya_adi']}",
                "info"
            )
            return path
        return None
    
    def prepare_output_dir(self, base_label, title):
        base_dir = filedialog.askdirectory(title=title)
        if not base_dir:
            return None
        folder_name = f"{base_label}_{datetime.now().strftime('%Y%m%d_%H%M')}"
        target_dir = os.path.join(base_dir, folder_name)
        os.makedirs(target_dir, exist_ok=True)
        return target_dir
    
    def copy_exam_files(self, base_dir, salon_names, create_subfolders=False):
        if not salon_names:
            return False
        success = False
        for salon_name in salon_names:
            salon = self.salon_data.get(salon_name)
            if not salon:
                continue
            target_dir = base_dir
            if create_subfolders:
                target_dir = os.path.join(base_dir, self._safe_folder_name(salon_name))
                os.makedirs(target_dir, exist_ok=True)
                # Şube/Salon kapak dosyası oluştur (liste başına gelecek şekilde isimlendir)
                self._create_cover_page(target_dir, salon_name)
                # Sınıf Oturma Düzeni Yoklama Listesi (Excel)
                self._create_seating_list_excel(target_dir, salon_name, salon)
                # Görsel Oturma Düzeni (Excel)
                self._create_visual_seating_excel(target_dir, salon_name, salon)
            if self._copy_exam_files_for_salon(target_dir, salon_name, salon):
                success = True
        return success
    
    def _create_seating_list_excel(self, target_dir, salon_name, salon):
        """Salon için sınıf oturma düzeni yoklama listesi Excel oluştur"""
        try:
            from controllers.excel_handler import ExcelHandler
            
            safe_name = self._safe_folder_name(salon_name)
            file_path = os.path.join(target_dir, f"000A_Yoklama_Listesi_{safe_name}.xlsx")
            
            sinav_bilgi = {
                'ders_adi': "SINAV",
                'tarih': datetime.now().strftime("%Y-%m-%d"),
                'saat': "-",
                'kacinci_ders': "-"
            }
            
            export_data = []
            for yer in salon.get('students', []):
                export_data.append({
                    'salon_adi': salon_name,
                    'sira_no': yer.get('sira_no', 0),
                    'ad': yer.get('ad', ''),
                    'soyad': yer.get('soyad', ''),
                    'sinif': yer.get('sinif', ''),
                    'sube': yer.get('sube', ''),
                    'gozetmenler': '',
                    'sinav_adi': yer.get('sinav_adi', '')
                })
            
            if export_data:
                ExcelHandler.yerlesim_yazdir(file_path, sinav_bilgi, export_data)
        except Exception as e:
            print(f"Yoklama listesi oluşturulamadı: {e}")
    
    def _create_visual_seating_excel(self, target_dir, salon_name, salon):
        """Salon için görsel oturma düzeni Excel oluştur"""
        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            
            safe_name = self._safe_folder_name(salon_name)
            file_path = os.path.join(target_dir, f"000B_Gorsel_Oturma_Duzeni_{safe_name}.xlsx")
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = salon_name[:30]
            
            # Sayfa Yapısı
            ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
            ws.page_setup.paperSize = ws.PAPERSIZE_A4
            ws.page_setup.fitToPage = True
            ws.page_setup.fitToHeight = 1
            ws.page_setup.fitToWidth = 1
            ws.page_margins.left = 0.3
            ws.page_margins.right = 0.3
            ws.page_margins.top = 0.3
            ws.page_margins.bottom = 0.3
            
            # Sütun Genişlikleri
            std_width = 18
            cor_width = 4
            for col, width in [('A', std_width), ('B', std_width), ('C', cor_width),
                               ('D', std_width), ('E', std_width), ('F', cor_width),
                               ('G', std_width), ('H', std_width)]:
                ws.column_dimensions[col].width = width
            
            # Stiller
            border_style = Border(
                left=Side(style='thin'), right=Side(style='thin'), 
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            red_font = Font(bold=True, size=11, color="FF0000")
            
            # Renk haritası
            CLASS_COLORS = {
                "5": "3498DB", "6": "2ECC71", "7": "E67E22", "8": "9B59B6",
                "9": "F1C40F", "10": "E74C3C", "11": "95A5A6", "12": "795548"
            }
            
            # Başlık
            ws.merge_cells('A1:H1')
            ws['A1'] = f"{salon_name} - SINAV OTURMA PLANI"
            ws['A1'].font = Font(bold=True, size=16)
            ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 45
            
            # KAPI, TAHTA, ÖĞRETMEN MASASI
            ws.row_dimensions[2].height = 35
            ws.merge_cells('A2:B2')
            ws['A2'] = "KAPI"
            ws['A2'].alignment = center_align
            ws['A2'].font = red_font
            ws['A2'].fill = PatternFill(start_color="FFEEEE", end_color="FFEEEE", fill_type="solid")
            
            ws.merge_cells('D2:E2')
            ws['D2'] = "TAHTA"
            ws['D2'].alignment = center_align
            ws['D2'].font = red_font
            ws['D2'].fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
            
            ws.merge_cells('G2:H2')
            ws['G2'] = "ÖĞRETMEN MASASI"
            ws['G2'].alignment = center_align
            ws['G2'].font = Font(bold=True, size=10)
            ws['G2'].fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
            
            # Koridor
            ws.merge_cells('C3:C7')
            ws['C3'] = "K\nO\nR\nİ\nD\nO\nR"
            ws['C3'].alignment = center_align
            ws['C3'].font = Font(size=8, color="888888")
            ws['C3'].fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
            
            ws.merge_cells('F3:F7')
            ws['F3'] = "K\nO\nR\nİ\nD\nO\nR"
            ws['F3'].alignment = center_align
            ws['F3'].font = Font(size=8, color="888888")
            ws['F3'].fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
            
            base_row_start = 3
            for r in range(base_row_start, base_row_start + 5):
                ws.row_dimensions[r].height = 75
            
            # Öğrenci map
            ogrenci_map = {}
            for yer in salon.get('students', []):
                try:
                    sn = int(yer['sira_no'])
                    ogrenci_map[sn] = yer
                except:
                    pass
            
            # 30 sıralık grid
            for sira_no in range(1, 31):
                block_idx = (sira_no - 1) // 10
                local_idx = (sira_no - 1) % 10
                row_in_block = local_idx // 2
                side = local_idx % 2
                
                if block_idx == 0:
                    final_row = base_row_start + row_in_block
                    final_col_idx = 1 if side == 0 else 2
                elif block_idx == 1:
                    final_row = (base_row_start + 4) - row_in_block
                    final_col_idx = 4 if side == 0 else 5
                else:
                    final_row = base_row_start + row_in_block
                    final_col_idx = 7 if side == 0 else 8
                
                col_letter = get_column_letter(final_col_idx)
                cell = ws[f"{col_letter}{final_row}"]
                
                ogr = ogrenci_map.get(sira_no)
                if ogr:
                    cell.value = f"{sira_no}\n{ogr.get('ad', '')} {ogr.get('soyad', '')}\n{ogr.get('sinif', '')}-{ogr.get('sube', '')}"
                    cell.alignment = center_align
                    cell.border = border_style
                    sinif_str = str(ogr.get('sinif', ''))
                    cls_key = "".join(filter(str.isdigit, sinif_str))
                    hex_color = CLASS_COLORS.get(cls_key, "BDC3C7")
                    cell.fill = PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")
                    cell.font = Font(size=9, color="FFFFFF")
                else:
                    cell.value = f"{sira_no}\n\nBOŞ"
                    cell.alignment = center_align
                    cell.border = border_style
                    cell.fill = PatternFill(start_color="EEEEEE", end_color="EEEEEE", fill_type="solid")
                    cell.font = Font(size=9, color="999999", italic=True)
            
            wb.save(file_path)
        except Exception as e:
            print(f"Görsel oturma düzeni oluşturulamadı: {e}")
    
    def _create_cover_page(self, target_dir, salon_name):
        """Salon/şube adını büyük puntolarla içeren kapak Word dosyası oluştur"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Sayfayı ayarla (dikey ortalama için boş alan)
            for _ in range(10):
                doc.add_paragraph("")
            
            # Salon adı - büyük puntolarla
            para = doc.add_paragraph()
            run = para.add_run(salon_name+"-SALONU")
            run.bold = True
            run.font.size = Pt(72)  # Çok büyük punto
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Alt başlık
            para2 = doc.add_paragraph()
            run2 = para2.add_run("SINAV KAĞITLARI")
            run2.bold = True
            run2.font.size = Pt(36)
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dosya adı - liste başına gelecek şekilde "000_" ile başlat
            safe_name = self._safe_folder_name(salon_name+"_SALONU")
            cover_path = os.path.join(target_dir, f"000_KAPAK_{safe_name}.docx")
            doc.save(cover_path)
        except Exception as e:
            print(f"Kapak sayfası oluşturulamadı: {e}")

    def _copy_exam_files_for_salon(self, target_dir, salon_name, salon):
        grouped = {}
        for yer in salon['students']:
            grouped.setdefault(yer.get('sinav_id'), []).append(yer)
        success = False
        for sinav_id, ogrenciler in grouped.items():
            exam = self.exam_lookup.get(sinav_id)
            source_file = self._ensure_exam_file(exam)
            if not source_file:
                continue
            ext = os.path.splitext(source_file)[1]
            for yer in sorted(ogrenciler, key=lambda x: x['sira_no']):
                seat_display = self._seat_display(yer)
                seat_tag = self._safe_folder_name(seat_display)
                exam_tag = self._safe_folder_name(yer.get('sinav_adi', "Sinav"))
                name_tag = self._safe_folder_name(f"{yer['ad']}_{yer['soyad']}")
                class_tag = self._safe_folder_name(f"{yer['sinif']}-{yer['sube']}")
                safe_name = f"{seat_tag}_{exam_tag}_{name_tag}_{class_tag}{ext}"
                dest = os.path.join(target_dir, safe_name)
                header_info = {
                    "student": f"{yer['ad']} {yer['soyad']}".strip(),
                    "class": f"{yer['sinif']}/{yer['sube']}",
                    "seat": seat_display,
                    "exam": yer.get('sinav_adi', exam.get('sinav_adi') if exam else "Sınav"),
                    "salon": salon_name
                }
                if ext.lower() == ".pdf":
                    if not self._personalize_pdf(source_file, dest, header_info):
                        shutil.copy2(source_file, dest)
                elif ext.lower() == ".docx":
                    if not self._personalize_docx(source_file, dest, header_info):
                        shutil.copy2(source_file, dest)
                else:
                    shutil.copy2(source_file, dest)
            success = True
        return success
    
    def print_selected_salon(self):
        if not self.selected_salon or self.selected_salon not in self.salon_data:
            show_message(self.window, "Lütfen salon seçin!", "warning")
            return
        target_dir = self.prepare_output_dir(
            self._safe_folder_name(self.selected_salon),
            "Seçili salon dosyalarının kopyalanacağı klasör"
        )
        if not target_dir:
            return
        if self.copy_exam_files(target_dir, [self.selected_salon]):
            show_message(self.window, f"📁 Dosyalar oluşturuldu:\n{target_dir}", "success")
    
    def print_all_available_salons(self):
        if not self.salon_data:
            show_message(self.window, "Yazdırılacak salon bulunamadı!", "warning")
            return
        
        base_dir = filedialog.askdirectory(title="Tüm salon dosyalarının kopyalanacağı klasör")
        if not base_dir:
            return
        root_dir = os.path.join(base_dir, f"Sinav_Kagitlari_tum_salonlar_{datetime.now().strftime('%Y%m%d_%H%M')}")
        os.makedirs(root_dir, exist_ok=True)
        
        if self.copy_exam_files(root_dir, sorted(self.salon_data.keys()), create_subfolders=True):
            show_message(self.window, f"📁 Sinav_Kagitlari_tum_salonlar oluşturuldu:\n{root_dir}", "success")
        else:
            show_message(self.window, "Yazdırılacak dosya oluşturulamadı!", "warning")

    def _safe_folder_name(self, text):
        if not text:
            return "sinav"
        allowed = []
        for ch in text:
            if ch.isalnum() or ch in (' ', '-', '_'):
                allowed.append(ch)
            else:
                allowed.append("_")
        return "".join(allowed).strip().replace(" ", "_") or "sinav"

    def _seat_display(self, yer: Dict) -> str:
        """Treeview'de gösterilecek sıra etiketi."""
        return format_sira_label(yer.get('sira_no'))

    def _personalize_pdf(self, source_path: str, dest_path: str, info: Dict[str, str]) -> bool:
        """PDF dosyasının üst kısmına öğrenci bilgisi ekle (tek satır)."""
        try:
            reader = PdfReader(source_path)
            writer = PdfWriter()
            # Tek satır format: İsim | Sınav | Salon | Sıra | Sınıf
            header_line = (
                f"{info.get('student', '-')}  |  "
                f"{info.get('exam', '-')}  |  "
                f"{info.get('salon', '-')}  |  "
                f"Sıra: {info.get('seat', '-')}  |  "
                f"{info.get('class', '-')}"
            )
            for page in reader.pages:
                width = float(page.mediabox.width)
                height = float(page.mediabox.height)
                packet = BytesIO()
                can = canvas.Canvas(packet, pagesize=(width, height))
                can.setFillColorRGB(0.12, 0.12, 0.12)
                can.setFont("Helvetica-Bold", 11)
                can.drawString(40, height - 25, header_line)
                can.line(35, height - 35, width - 35, height - 35)
                can.save()
                packet.seek(0)
                overlay = PdfReader(packet).pages[0]
                page.merge_page(overlay)
                writer.add_page(page)
            with open(dest_path, "wb") as target:
                writer.write(target)
            return True
        except Exception as exc:
            print(f"PDF kişiselleştirme başarısız ({source_path}): {exc}")
            return False

    def _personalize_docx(self, source_path: str, dest_path: str, info: Dict[str, str]) -> bool:
        """DOCX dosyasının üst bilgisini öğrenci adları ile doldur (tek satır)."""
        try:
            doc = Document(source_path)
            if not doc.sections:
                doc.add_section()
            section = doc.sections[0]
            header = section.header
            while len(header.paragraphs) > 0:
                p = header.paragraphs[0]
                p._p.getparent().remove(p._p)
            # Tek satır format: İsim | Sınav | Salon | Sıra | Sınıf
            header_para = header.add_paragraph()
            header_para.text = (
                f"{info.get('student', '-')}  |  "
                f"Sınav:{info.get('exam', '-')}  |  "
                f"Salon:{info.get('salon', '-')}  |  "
                f"Sıra: {info.get('seat', '-')}  |  "
                
            )
            header_para.style = header_para.style or doc.styles['Normal']
            header_para.runs[0].font.bold = True
            doc.save(dest_path)
            return True
        except Exception as exc:
            print(f"DOCX kişiselleştirme başarısız ({source_path}): {exc}")
            return False
